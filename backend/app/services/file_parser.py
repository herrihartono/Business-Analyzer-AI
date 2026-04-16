from __future__ import annotations

import os
from pathlib import Path

import polars as pl
import pdfplumber
from docx import Document

from app.config import get_settings

settings = get_settings()


def parse_file(stored_filename: str) -> pl.DataFrame:
    """Parse an uploaded file into a Polars DataFrame."""
    path = os.path.join(settings.upload_dir, stored_filename)
    ext = Path(stored_filename).suffix.lower()

    if ext == ".csv":
        return _parse_csv(path)
    elif ext in (".xlsx", ".xls"):
        return _parse_excel(path)
    elif ext == ".pdf":
        return _parse_pdf(path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_csv(path: str) -> pl.DataFrame:
    return pl.read_csv(path, infer_schema_length=1000, try_parse_dates=True)


def _parse_excel(path: str) -> pl.DataFrame:
    return pl.read_excel(path, infer_schema_length=1000)


def _parse_pdf(path: str) -> pl.DataFrame:
    """Extract tables from PDF. Falls back to extracting text as rows."""
    tables: list[list[list[str]]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            if page_tables:
                tables.extend(page_tables)

    if tables:
        merged = tables[0]
        for t in tables[1:]:
            merged.extend(t)
        if len(merged) < 2:
            return pl.DataFrame({"text": [str(row) for row in merged]})
        headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(merged[0])]
        data = {h: [] for h in headers}
        for row in merged[1:]:
            for i, h in enumerate(headers):
                data[h].append(row[i] if i < len(row) else None)
        return pl.DataFrame(data)

    with pdfplumber.open(path) as pdf:
        lines = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.extend(text.split("\n"))
    return pl.DataFrame({"text": lines}) if lines else pl.DataFrame({"text": ["(empty document)"]})


def _parse_docx(path: str) -> pl.DataFrame:
    """Extract tables from DOCX. Falls back to paragraph text."""
    doc = Document(path)

    if doc.tables:
        table = doc.tables[0]
        headers = [cell.text.strip() or f"col_{i}" for i, cell in enumerate(table.rows[0].cells)]
        data = {h: [] for h in headers}
        for row in table.rows[1:]:
            for i, h in enumerate(headers):
                data[h].append(row.cells[i].text.strip() if i < len(row.cells) else None)
        return pl.DataFrame(data)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return (
        pl.DataFrame({"text": paragraphs})
        if paragraphs
        else pl.DataFrame({"text": ["(empty document)"]})
    )


def dataframe_preview(df: pl.DataFrame, n: int = 50) -> list[dict]:
    """Return first n rows as list of dicts for JSON serialization."""
    return df.head(n).to_dicts()


def column_statistics(df: pl.DataFrame) -> dict:
    """Compute basic statistics for each column."""
    stats: dict = {}
    for col in df.columns:
        dtype = str(df[col].dtype)
        info: dict = {"dtype": dtype, "null_count": df[col].null_count(), "unique_count": df[col].n_unique()}

        if df[col].dtype in (pl.Float64, pl.Float32, pl.Int64, pl.Int32, pl.Int16, pl.Int8):
            desc = df[col].drop_nulls()
            if len(desc) > 0:
                info["mean"] = float(desc.mean())  # type: ignore[arg-type]
                info["min"] = float(desc.min())  # type: ignore[arg-type]
                info["max"] = float(desc.max())  # type: ignore[arg-type]
                info["std"] = float(desc.std())  # type: ignore[arg-type]

        stats[col] = info
    return stats
